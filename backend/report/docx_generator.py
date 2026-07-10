import json

from docx import Document


def generate_docx(job, aggregates: dict, output_path: str) -> None:
    doc = Document()
    doc.add_heading("Báo cáo NGS Monitor", level=1)
    doc.add_paragraph(f"Khoảng thời gian: {job.date_from} – {job.date_to}")

    doc.add_heading("Bảng 3.1. Số lượng nội dung theo cơ quan", level=2)
    _add_count_table(doc, "Cơ quan", aggregates["source_counts"])

    doc.add_heading("Bảng 3.2 / 3.7. Số lượng nội dung theo chủ đề", level=2)
    _add_count_table(doc, "Chủ đề", aggregates["topic_counts"])

    doc.add_heading("Bảng 3.8. Top từ khóa", level=2)
    _add_count_table(doc, "Từ khóa", aggregates["keyword_counts"])

    doc.add_heading("Thống kê số lượng nội dung theo tháng (tương ứng Hình 3.2 gốc, dạng bảng)", level=2)
    _add_count_table(doc, "Tháng", aggregates["monthly_counts"])

    doc.add_heading("Bảng 3.13. Kết quả phân tích sắc thái cảm xúc (Sentiment Analysis)", level=2)
    _add_count_table(doc, "Sentiment", aggregates["sentiment_counts"])

    doc.add_heading("Bảng 3.15. Kết quả phân tích cảm xúc (Emotion Analysis)", level=2)
    _add_count_table(doc, "Emotion", aggregates["emotion_counts"])

    doc.add_heading("Bảng 3.17. Thống kê tổng hợp", level=2)
    _add_count_table(doc, "Chỉ số", aggregates["summary_stats"])

    doc.add_heading("Danh sách bài viết", level=2)
    table = doc.add_table(rows=1, cols=6)
    header = table.rows[0].cells
    for i, name in enumerate(["Tiêu đề", "URL", "Sentiment", "Emotion", "Confidence", "Needs review"]):
        header[i].text = name
    for article in aggregates["articles"]:
        cells = table.add_row().cells
        cells[0].text = article["title"] or ""
        cells[1].text = article["url"] or ""
        cells[2].text = article["sentiment"] or ""
        cells[3].text = article["emotion"] or ""
        cells[4].text = str(article["confidence"])
        cells[5].text = str(article["needs_review"])

    doc.save(output_path)


def _add_count_table(doc: Document, label_column: str, counts: dict) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = label_column
    table.rows[0].cells[1].text = "Số lượng"
    for key, count in counts.items():
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(count)


def export_json(job, aggregates: dict, output_path: str) -> None:
    data = {
        "job_id": str(job.job_id),
        "date_from": str(job.date_from),
        "date_to": str(job.date_to),
        **aggregates,
    }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2, default=str)
