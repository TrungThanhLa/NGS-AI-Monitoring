import json
import uuid
from datetime import date, datetime
from types import SimpleNamespace

import docx

from backend.report.docx_generator import export_json, generate_docx

AGGREGATES = {
    "articles": [
        {
            "title": "Cảnh báo deepfake mới",
            "url": "https://vtv.vn/bai-1",
            "source": "VTV News",
            "published_at": datetime(2026, 6, 23, 10, 0, 0),
            "sentiment": "negative",
            "emotion": "Fear",
            "topics": ["AI, Deepfake và công nghệ tạo sinh"],
            "confidence": 0.9,
            "needs_review": False,
            "summary": "Tóm tắt bài viết.",
        }
    ],
    "sentiment_counts": {"negative": 1},
    "emotion_counts": {"Fear": 1},
    "source_counts": {"VTV": 1},
    "topic_counts": {"AI, Deepfake và công nghệ tạo sinh": 1},
    "keyword_counts": {"deepfake": 1},
    "monthly_counts": {"2026-06": 1},
    "summary_stats": {"Tổng số bài": 1, "Tổng số cơ quan": 1, "Số bài cần review (needs_review)": 0},
}


def _fake_job():
    return SimpleNamespace(job_id=uuid.uuid4(), date_from=date(2026, 6, 1), date_to=date(2026, 6, 30))


def test_generate_docx_writes_readable_file_with_article_title(tmp_path):
    output_path = tmp_path / "report.docx"

    generate_docx(_fake_job(), AGGREGATES, str(output_path))

    assert output_path.exists()
    doc = docx.Document(str(output_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "Cảnh báo deepfake mới" in table_text or "Cảnh báo deepfake mới" in full_text


def test_export_json_writes_valid_json_with_counts(tmp_path):
    output_path = tmp_path / "report.json"
    job = _fake_job()

    export_json(job, AGGREGATES, str(output_path))

    data = json.loads(output_path.read_text(encoding="utf-8"))
    assert data["sentiment_counts"] == {"negative": 1}
    assert data["emotion_counts"] == {"Fear": 1}
    assert data["articles"][0]["title"] == "Cảnh báo deepfake mới"
    assert data["job_id"] == str(job.job_id)


def test_generate_docx_includes_new_aggregate_tables(tmp_path):
    output_path = tmp_path / "report.docx"

    generate_docx(_fake_job(), AGGREGATES, str(output_path))

    doc = docx.Document(str(output_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    combined = full_text + "\n" + table_text

    assert "Bảng 3.1" in combined
    assert "VTV" in table_text
    assert "AI, Deepfake và công nghệ tạo sinh" in table_text
    assert "deepfake" in table_text
    assert "2026-06" in table_text
    assert "Tổng số bài" in table_text


def test_article_list_appears_after_aggregate_tables(tmp_path):
    output_path = tmp_path / "report.docx"

    generate_docx(_fake_job(), AGGREGATES, str(output_path))

    doc = docx.Document(str(output_path))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

    assert headings.index("Bảng 3.1. Số lượng nội dung theo cơ quan") < headings.index("Danh sách bài viết")


def test_export_json_includes_new_aggregate_fields(tmp_path):
    output_path = tmp_path / "report.json"
    job = _fake_job()

    export_json(job, AGGREGATES, str(output_path))

    data = json.loads(output_path.read_text(encoding="utf-8"))
    assert data["source_counts"] == {"VTV": 1}
    assert data["topic_counts"] == {"AI, Deepfake và công nghệ tạo sinh": 1}
    assert data["keyword_counts"] == {"deepfake": 1}
    assert data["monthly_counts"] == {"2026-06": 1}
    assert data["summary_stats"]["Tổng số bài"] == 1
